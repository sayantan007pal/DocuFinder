"""
src/extraction/local_sheets.py — LocalSheetExtractor.
Primary: PyMuPDF (fitz) page.find_tables() — fast, simple tables.
DOCX: python-docx (always high confidence — structured format).
"""
import time
from pathlib import Path

import structlog

from src.extraction.sheets import (
    BaseSheetExtractor,
    ExtractedTable,
    ExtractionResult,
    TableConfidence,
)

log = structlog.get_logger(__name__)


class LocalSheetExtractor(BaseSheetExtractor):
    """
    Local table extraction using PyMuPDF.

    Strategy:
    1. Use PyMuPDF page.find_tables() — fast, works for simple tables
    2. For DOCX files, use python-docx directly
    """

    def __init__(self):
        pass  # No external dependencies needed

    async def extract_tables(
        self,
        file_path: str,
        doc_id: str,
        tenant_id: str,
        page_numbers: list[int] | None = None,
    ) -> ExtractionResult:
        t0 = time.monotonic()
        path = Path(file_path)

        if path.suffix.lower() == ".pdf":
            tables = await self._extract_pdf_tables(file_path, page_numbers)
            provider = "local_pymupdf"
        elif path.suffix.lower() in (".docx", ".doc"):
            tables = await self._extract_docx_tables(file_path)
            provider = "python_docx"
        else:
            raise ValueError(f"Unsupported file type: {path.suffix}")

        return ExtractionResult(
            doc_id=doc_id,
            tenant_id=tenant_id,
            filename=path.name,
            tables=tables,
            extraction_time_ms=(time.monotonic() - t0) * 1000,
            provider_used=provider,
        )

    async def _extract_pdf_tables(
        self,
        file_path: str,
        page_numbers: list[int] | None,
    ) -> list[ExtractedTable]:
        import fitz

        tables = []
        doc = fitz.open(file_path)
        pages_to_process = page_numbers or list(range(len(doc)))

        for page_idx in pages_to_process:
            if page_idx >= len(doc):
                continue
            page = doc[page_idx]

            try:
                page_tables_obj = page.find_tables()
                page_tables = page_tables_obj.tables if page_tables_obj else []
            except AttributeError:
                # Older PyMuPDF versions don't have find_tables
                page_tables = []

            for table_idx, table in enumerate(page_tables):
                data = table.extract()
                if not data:
                    continue

                headers = [str(h) if h else "" for h in (data[0] if data else [])]
                rows = [[str(c) if c else "" for c in row] for row in (data[1:] if len(data) > 1 else [])]
                confidence = self._assess_confidence(headers, rows, table)

                tables.append(ExtractedTable(
                    page_number=page_idx + 1,
                    table_index=table_idx,
                    headers=headers,
                    rows=rows,
                    bbox=tuple(table.bbox) if hasattr(table, "bbox") else None,
                    confidence=confidence,
                    source_parser="pymupdf",
                ))

        doc.close()
        return tables

    async def _extract_docx_tables(self, file_path: str) -> list[ExtractedTable]:
        """Extract tables from DOCX using python-docx."""
        from docx import Document as DocxDocument

        tables = []
        doc = DocxDocument(file_path)

        for table_idx, table in enumerate(doc.tables):
            rows_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                rows_data.append(row_data)

            if not rows_data:
                continue

            headers = rows_data[0] if rows_data else []
            rows = rows_data[1:] if len(rows_data) > 1 else []

            tables.append(ExtractedTable(
                page_number=1,
                table_index=table_idx,
                headers=headers,
                rows=rows,
                confidence=TableConfidence.HIGH,
                source_parser="python_docx",
            ))

        return tables

    def _assess_confidence(self, headers, rows, table) -> TableConfidence:
        """Assess extraction confidence based on table structure."""
        if not headers or not rows:
            return TableConfidence.LOW

        header_count = len(headers)
        consistent_cols = all(len(row) == header_count for row in rows)
        empty_cells = sum(1 for row in rows for cell in row if not str(cell).strip())
        total_cells = sum(len(row) for row in rows)
        empty_ratio = empty_cells / max(total_cells, 1)

        if consistent_cols and empty_ratio < 0.1:
            return TableConfidence.HIGH
        elif consistent_cols or empty_ratio < 0.3:
            return TableConfidence.MEDIUM
        else:
            return TableConfidence.LOW

    async def health_check(self) -> bool:
        return True  # Local extraction always available
