"""
src/ingestion/parsers/pythondocx.py — Pure Python DOCX parser.

Uses python-docx library — no LibreOffice dependency.
Only supports .docx files (not .doc legacy format).

Usage: Set PARSER_DOCX=python-docx in .env
"""
from pathlib import Path

import structlog
from llama_index.core.schema import Document

from src.ingestion.parsers.base import BaseDocumentParser

log = structlog.get_logger(__name__)


class PythonDocxParser(BaseDocumentParser):
    """
    Pure Python DOCX parser using python-docx library.
    No external dependencies like LibreOffice required.
    
    Usage: Set PARSER_DOCX=python-docx in .env
    """

    async def parse(
        self,
        file_path: str,
        doc_id: str,
        tenant_id: str,
    ) -> list[Document]:
        from docx import Document as DocxDocument
        from docx.opc.exceptions import PackageNotFoundError

        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix not in (".docx",):
            raise ValueError(
                f"PythonDocxParser only supports .docx files, got: {suffix}. "
                "Use PARSER_PROVIDER=liteparse or unstructured for other formats."
            )

        log.info("pythondocx_start", file=path.name, doc_id=doc_id)

        try:
            docx = DocxDocument(file_path)
        except PackageNotFoundError as exc:
            log.error("pythondocx_invalid_file", file=path.name, error=str(exc))
            raise RuntimeError(f"Invalid DOCX file: {path.name}") from exc
        except Exception as exc:
            log.error("pythondocx_parse_error", file=path.name, error=str(exc))
            raise RuntimeError(f"Failed to parse DOCX: {exc}") from exc

        base_meta = self._base_metadata(file_path, doc_id, tenant_id, "python-docx")

        documents: list[Document] = []

        # Extract text from paragraphs
        paragraphs_text = []
        for para in docx.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs_text.append(text)

        # Extract text from tables
        tables_text = []
        for table in docx.tables:
            table_rows = []
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    table_rows.append(row_text)
            if table_rows:
                tables_text.append("\n".join(table_rows))

        # Combine all text
        all_text_parts = paragraphs_text + tables_text
        full_text = "\n\n".join(all_text_parts)

        if full_text.strip():
            # DOCX doesn't have pages, so we create a single document
            # with page_number=1 for consistency
            documents.append(Document(
                text=full_text,
                metadata={**base_meta, "page_number": 1},
            ))

        log.info("pythondocx_done", file=path.name, doc_count=len(documents),
                 paragraphs=len(paragraphs_text), tables=len(tables_text))
        return documents
